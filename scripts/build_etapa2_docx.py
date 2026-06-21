"""Gera o docx da Etapa 2 — modelagem causal via Double ML.

OBSOLETO: este script tem conteúdo da versão de 3 temporadas (2.079 obs, econml)
hardcoded. A fonte da verdade agora é `docs/etapa2_modelagem_causal.md` (8 temporadas,
DML manual). Para regenerar o docx atual, use pandoc:
    pandoc docs/etapa2_modelagem_causal.md -o docs/etapa2_modelagem_causal.docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PRIMARY = RGBColor(0x0B, 0x3D, 0x91)
ACCENT  = RGBColor(0xE2, 0x7D, 0x60)
DARK    = RGBColor(0x1F, 0x2A, 0x44)
MUTED   = RGBColor(0x6B, 0x72, 0x80)

doc = Document()

# ---------- estilos base ----------
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = DARK

# margens
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def h(level, text, color=PRIMARY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt({1: 18, 2: 14, 3: 12}.get(level, 11))


def para(text, *, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def para_mixed(parts):
    """parts: list of (text, {'bold':..., 'italic':..., 'mono':...})."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    for txt, opts in parts:
        r = p.add_run(txt)
        r.bold = opts.get("bold", False)
        r.italic = opts.get("italic", False)
        if opts.get("mono"):
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        else:
            r.font.size = Pt(opts.get("size", 11))
    return p


def quote_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    # barra lateral cinza
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "E27D60")
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = MUTED
    run.font.size = Pt(10.5)


def equation(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Cambria Math"
    run.font.size = Pt(12)
    run.italic = True
    run.font.color.rgb = DARK


def bullet(items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(it)
        run.font.size = Pt(11)


def code_block(code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F1F5F9")
    pPr.append(shd)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = DARK


def add_table(header, rows, *, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
    tbl.style = "Light Grid Accent 1"
    tbl.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            for r in tbl.rows:
                r.cells[i].width = w
    # header
    for i, txt in enumerate(header):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        set_cell_shading(cell, "0B3D91")
        p = cell.paragraphs[0]
        r = p.add_run(txt)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.name = "Calibri"
    # rows
    for i, row in enumerate(rows):
        for j, txt in enumerate(row):
            cell = tbl.rows[i + 1].cells[j]
            cell.text = ""
            if i % 2 == 1:
                set_cell_shading(cell, "F1F5F9")
            p = cell.paragraphs[0]
            r = p.add_run(str(txt))
            r.font.size = Pt(10)
            r.font.name = "Calibri"
            r.font.color.rgb = DARK
    # margem zero entre tabela e parágrafo seguinte
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return tbl


# ===================== TÍTULO =====================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
trun = title.add_run("Etapa 2 — Estimação Causal por Aprendizado de Máquina Duplo")
trun.bold = True
trun.font.size = Pt(20)
trun.font.color.rgb = PRIMARY

quote_block(
    "Versão adaptada à realidade do dataset transfers_etapa2_ready.csv "
    "(2.079 transferências · 3 temporadas · 7 ligas). Substitui a Etapa 2 "
    "descrita no documento de estratégia original onde a proposta dependia "
    "de variáveis que ainda não estão disponíveis na base (data exata da "
    "transferência, contexto esportivo, performance individual)."
)

# ===================== 1. DA ETAPA 1 PARA A ETAPA 2 =====================
h(2, "1. Da Etapa 1 para a Etapa 2")

para("A Etapa 1 entregou o resíduo hedônico de cada transferência:")
equation("Y(i,c) = ln(P_i) − ln(P̂_i)")
para(
    "onde P̂_i é o preço justo previsto pelo modelo Random Forest treinado "
    "apenas com atributos do jogador e controles estruturais de mercado "
    "(idade, idade², log do market value, posição, dummies de liga e temporada)."
)
para(
    "Por construção, Y(i,c) está livre de características intrínsecas do "
    "atleta. O que sobra no resíduo é, por hipótese, forças conjunturais "
    "da negociação. A Etapa 2 testa se uma dessas forças é a liquidez "
    "recente do comprador."
)

# ===================== 2. TRATAMENTO =====================
h(2, "2. Definição operacional de tratamento")

para(
    "A proposta original definia o tratamento como o logaritmo do volume "
    "acumulado de vendas do clube nos 30 dias anteriores à compra. A base "
    "disponível, porém, agrega transações por temporada, não por dia. "
    "Operacionalizamos o tratamento na granularidade efetivamente disponível:"
)
equation("D_c = ln(1 + revenue_sales_{c,t})")
para(
    "onde revenue_sales_{c,t} é a soma de fees recebidos pelo clube c em "
    "vendas durante a temporada t. Variantes binárias (big_sale_flag = 1 "
    "se max_sale_c > €30M) podem ser testadas como robustez para isolar "
    "o efeito de vendas blockbuster."
)

# ===================== 3. MODELO =====================
h(2, "3. Especificação do modelo Double ML")

para("Seguimos o modelo parcialmente linear de Chernozhukov et al. (2018):")
equation("Y = θ·D + g(W) + U,   E[U | W, D] = 0")
equation("D = m(W) + V,         E[V | W] = 0")
para(
    "onde W é o vetor de confundidores estruturais do comprador, da liga "
    "e da temporada. A escolha de W é direta dado o que o dataset oferece "
    "e está detalhada na Tabela 1."
)

h(3, "Tabela 1 — Mapeamento de W aos confundidores neutralizados")
add_table(
    header=["Bloco", "Variáveis em W", "Confundidor neutralizado"],
    rows=[
        ["Financeiro do clube",
         "total_spend · n_buys · net_balance · net_transfer_record",
         "C1 — poder financeiro"],
        ["Estrutural do elenco",
         "squad_size · average_age · national_team_players",
         "C1, C6 — perfil/prestígio"],
        ["Posição na rede",
         "pagerank · in_strength · out_strength · net_flow · in_degree · out_degree",
         "C1, C6 — poder de barganha estrutural"],
        ["Mercado / liga",
         "log_league_mv · 6 dummies competition_code_*",
         "C4, C5 — sazonalidade e inflação"],
        ["Temporada",
         "season_2024 · season_2025",
         "C4 — efeito de calendário"],
    ],
    col_widths=[Cm(3.2), Cm(8.5), Cm(5.0)],
)

para(
    "A inclusão das features de rede é uma melhoria sobre a proposta original, "
    "que tratava apenas de receita, dias restantes na janela, team rating e "
    "classificação para a UCL — três das quatro não estão disponíveis na base "
    "atual. O PageRank e as métricas de força capturam de forma compacta o "
    "mesmo sinal de “poder estrutural” que aquelas variáveis pretendiam "
    "representar."
)

h(3, "Confundidores ainda não controlados")
add_table(
    header=["ID", "Bias", "Estratégia futura"],
    rows=[
        ["C2", "Causalidade reversa (compra precede venda, multa rescisória)",
         "Enriquecer com flag de multa via scraping do Transfermarkt"],
        ["C3", "Causa comum (classificação para UCL, troca de técnico)",
         "Integrar dataset de campeonatos e calendário esportivo"],
    ],
    col_widths=[Cm(1.5), Cm(7.5), Cm(7.7)],
)
para(
    "Esses dois confundidores permanecem como limitação reconhecida do desenho "
    "atual e devem ser abordados na próxima iteração da modelagem."
)

# ===================== 4. PROTOCOLO =====================
h(2, "4. Protocolo algorítmico")

para("A estimação segue os três pilares do DML:")
bullet([
    "Ortogonalização de Neyman. O parâmetro causal θ é obtido pela regressão "
    "linear simples do resíduo Ỹ = Y − ĝ(W) contra o resíduo D̃ = D − m̂(W). "
    "A propriedade de ortogonalidade garante que erros de primeira ordem em "
    "ĝ e m̂ não viciem θ̂.",
    "Cross-fitting. A amostra é dividida em K = 5 partições; ĝ e m̂ são treinados "
    "em K−1 partições e os resíduos são avaliados na partição restante, em "
    "rodízio. Isso elimina o viés de sobreajuste típico do reuso de dados.",
    "Inferência robusta. Erros-padrão são clusterizados por clube comprador "
    "para acomodar correlação intra-clube ao longo das temporadas.",
])

h(3, "Implementação em Python (EconML)")
code_block(
'''import numpy as np
import pandas as pd
from econml.dml import LinearDML, CausalForestDML
from sklearn.ensemble import RandomForestRegressor

df = pd.read_csv("output/transfers_etapa2_ready.csv")

Y = df["premio_reinvestimento"].values
D = df["log_revenue"].values

W_cols = [
    # financeiro
    "total_spend", "n_buys", "net_balance", "net_transfer_record",
    # elenco
    "squad_size", "average_age", "national_team_players",
    # rede
    "pagerank", "in_strength", "out_strength", "net_flow",
    "in_degree", "out_degree",
    # liga e temporada
    "log_league_mv", "season_2024", "season_2025",
    "competition_code_premier-league", "competition_code_laliga",
    "competition_code_serie-a", "competition_code_ligue-1",
    "competition_code_liga-portugal", "competition_code_jupiler-pro-league",
]
W = df[W_cols].fillna(0).values

dml = LinearDML(
    model_y=RandomForestRegressor(n_estimators=300, max_depth=5, random_state=42),
    model_t=RandomForestRegressor(n_estimators=300, max_depth=5, random_state=42),
    discrete_treatment=False,
    cv=5,
    random_state=42,
)
dml.fit(Y, D, X=None, W=W, inference="auto")
print(dml.summary())'''
)
para(
    "A saída fornece θ̂, erro-padrão, intervalo de confiança 95% e p-valor "
    "para a hipótese H₀: θ = 0 (ausência de prêmio do vendedor)."
)

# ===================== 5. MÉTRICAS INOVADORAS =====================
h(2, "5. Métricas inovadoras adaptadas")
para(
    "A proposta original sugeria duas métricas: a Curva de Decaimento "
    "Temporal e o Índice de Vulnerabilidade de Barganha. A primeira é "
    "inviável na base atual porque exige a data exata da transferência. "
    "Mantemos a segunda integralmente e propomos uma terceira métrica, "
    "originalmente nossa, que aproveita a disponibilidade das features "
    "de rede."
)

h(3, "5.1 Índice de Vulnerabilidade de Barganha (IVB)")
para(
    "Estimando efeitos heterogêneos com CausalForestDML, obtemos θ_c para "
    "cada clube comprador presente em pelo menos duas temporadas. O IVB "
    "normaliza esse efeito no intervalo [0, 1]:"
)
equation("IVB_c = (θ_c − min(Θ)) / (max(Θ) − min(Θ))")
para(
    "Clubes com IVB alto são “presas fáceis” — pagam sobrepreços maiores "
    "quando chegam ao mercado recém-capitalizados. Clubes com IVB baixo "
    "são negociadores disciplinados, capazes de reciclar capital sem "
    "ceder à pressão."
)
para(
    "A mesma lógica, aplicada do lado vendedor, identifica os “clubes "
    "predadores” — times que extraem as maiores taxas de prêmio positivo "
    "quando abordados por compradores recém-capitalizados."
)

h(3, "5.2 Sensibilidade do prêmio à centralidade na rede (substitui θ(Δt))")
para(
    "Sem datas diárias, não podemos modelar o decaimento temporal proposto. "
    "Em contrapartida, exploramos uma dimensão ortogonal que a base permite: "
    "a heterogeneidade do efeito segundo a posição estrutural do comprador "
    "na rede de transferências. Parametrizamos:"
)
equation("θ(PageRank_c) = θ_0 + θ_1 · quartil(PageRank_c)")
para(
    "A hipótese testada é que clubes centrais (alta centralidade de "
    "PageRank, top 25% do mercado) sofrem prêmios menores que clubes "
    "periféricos, por deterem informação superior e poder de barganha "
    "estrutural. Essa parametrização entrega um insight gerencial direto: "
    "quanto vale, em euros, estar posicionado no núcleo da rede de "
    "transferências."
)

h(3, "5.3 Curva temporal por temporada (proxy do decaimento)")
para(
    "Como degradação aceitável da curva θ(Δt) original, estimamos θ_t "
    "separadamente para cada temporada t ∈ {2023, 2024, 2025}. A "
    "tendência θ_{2023} → θ_{2025} indica se o prêmio do vendedor se "
    "intensificou ou se dissipou com a correção de mercado observada "
    "em 2025."
)

# ===================== 6. ROBUSTEZ =====================
h(2, "6. Testes de robustez planejados")
bullet([
    "Especificação de D. Replicar com D = log_revenue, com n_sales e com "
    "a flag binária big_sale.",
    "Subgrupos. Estimar θ por tier de comprador (top 5 ligas vs. demais) "
    "e por posição do jogador.",
    "Placebo. Aleatorizar D entre clubes mantendo a estrutura de W; o θ "
    "estimado deve ser estatisticamente indistinguível de zero.",
    "Reverse causation. Reestimar com tratamento defasado (lag de uma "
    "temporada) — se o efeito persistir, a hipótese de causalidade reversa "
    "(C2) enfraquece.",
    "PSM como sanity check do C6. Construir pares de clubes com propensity "
    "score similar (alto vs. baixo log_revenue) e comparar o prêmio médio "
    "diretamente, sem o ferramental DML.",
])

# ===================== 7. CRONOGRAMA =====================
h(2, "7. Cronograma de implementação")
add_table(
    header=["Semana", "Entrega"],
    rows=[
        ["1", "Pipeline DML completo (LinearDML + CausalForestDML), "
              "análises de robustez 1 e 3"],
        ["2", "Cálculo de IVB por clube; ranking de clubes-presa e "
              "clubes-predadores"],
        ["2", "Curva θ_t por temporada e sensibilidade ao PageRank"],
        ["3", "Validação por PSM e roadmap de enriquecimento "
              "(data exata, performance)"],
    ],
    col_widths=[Cm(2.2), Cm(14.5)],
)

# ===================== 8. LIMITAÇÕES =====================
h(2, "8. Limitações reconhecidas")
para(
    "A modelagem atual não controla por contexto esportivo individual "
    "(classificação para a UCL, troca de técnico, choques de receita "
    "de TV) nem isola transferências motivadas por cláusula de multa "
    "rescisória. Esses elementos correspondem aos confundidores C2 e "
    "C3 e ficam endereçados como trabalho futuro. A granularidade "
    "temporal trimestral (por temporada) é a restrição mais relevante: "
    "ela limita a precisão da estimativa de urgência e impede a "
    "estimação direta da curva exponencial de decaimento. Ainda assim, "
    "o desenho proposto é suficiente para responder à pergunta principal "
    "do projeto — existe um prêmio causal do vendedor? — com rigor "
    "metodológico compatível com o estado da arte."
)

# ---------- salva ----------
out = "/Users/lucaspedras/TP-CDAF/docs/etapa2_modelagem_causal.docx"
doc.save(out)
print(f"OK: {out}")
