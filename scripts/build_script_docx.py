"""Gera o docx do roteiro de apresentação.

OBSOLETO: este script tem conteúdo da versão de 3 temporadas (2.079 obs) hardcoded.
A fonte da verdade agora é `docs/script_apresentacao.md`. Para regenerar o docx atual:
    pandoc docs/script_apresentacao.md -o docs/script_apresentacao.docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PRIMARY = RGBColor(0x0B, 0x3D, 0x91)
ACCENT  = RGBColor(0xE2, 0x7D, 0x60)
DARK    = RGBColor(0x1F, 0x2A, 0x44)
MUTED   = RGBColor(0x6B, 0x72, 0x80)
GREEN   = RGBColor(0x2E, 0x7D, 0x32)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = DARK

for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_para_shading(p, hex_color):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def set_left_border(p, color_hex="E27D60", size="18"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)


def h(level, text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = color or PRIMARY
    r.font.size = Pt({1: 20, 2: 15, 3: 12}.get(level, 11))


def slide_header(num, title, duration):
    """Cabeçalho destacado de cada slide do roteiro."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    set_para_shading(p, "0B3D91")

    r1 = p.add_run(f"  Slide {num}  ")
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r1.bold = True; r1.font.size = Pt(13)

    r2 = p.add_run(f"·  {title}  ")
    r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r2.font.size = Pt(13)

    r3 = p.add_run(f"·  ⏱ {duration}  ")
    r3.font.color.rgb = RGBColor(0xE2, 0x7D, 0x60)
    r3.bold = True
    r3.font.size = Pt(13)


def script_para(parts):
    """Parágrafo de fala. parts: lista de (texto, opts)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.35
    set_left_border(p, color_hex="0B3D91", size="14")
    for txt, opts in parts:
        r = p.add_run(txt)
        r.font.size = Pt(11.5)
        r.bold = opts.get("bold", False)
        r.italic = opts.get("italic", False)
        if opts.get("muted"):
            r.font.color.rgb = MUTED
        elif opts.get("accent"):
            r.font.color.rgb = ACCENT
            r.bold = True
        else:
            r.font.color.rgb = DARK
    return p


def cue(text):
    """Marcador [clica] / [gesto] em itálico cinza."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.italic = True
    r.font.color.rgb = MUTED
    r.font.size = Pt(10)


def para(text, *, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def add_table(header, rows, *, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
    tbl.style = "Light Grid Accent 1"
    tbl.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            for r in tbl.rows:
                r.cells[i].width = w
    for i, txt in enumerate(header):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        set_cell_shading(cell, "0B3D91")
        p = cell.paragraphs[0]
        r = p.add_run(txt)
        r.bold = True; r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.name = "Calibri"
    for i, row in enumerate(rows):
        for j, txt in enumerate(row):
            cell = tbl.rows[i + 1].cells[j]
            cell.text = ""
            if i % 2 == 1:
                set_cell_shading(cell, "F1F5F9")
            p = cell.paragraphs[0]
            r = p.add_run(str(txt))
            r.font.size = Pt(10); r.font.name = "Calibri"
            r.font.color.rgb = DARK
    return tbl


def qa(question, answer_parts):
    """Bloco de Q&A: pergunta destacada + resposta."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run("P: ")
    r1.bold = True; r1.font.color.rgb = ACCENT; r1.font.size = Pt(11)
    r2 = p.add_run(question)
    r2.italic = True; r2.font.size = Pt(11); r2.font.color.rgb = DARK

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.6)
    p2.paragraph_format.space_after = Pt(8)
    p2.paragraph_format.line_spacing = 1.3
    rR = p2.add_run("R: ")
    rR.bold = True; rR.font.color.rgb = GREEN; rR.font.size = Pt(11)
    for txt, opts in answer_parts:
        r = p2.add_run(txt)
        r.font.size = Pt(11)
        r.bold = opts.get("bold", False)
        r.italic = opts.get("italic", False)
        r.font.color.rgb = DARK


# ============================== CAPA ==============================
title = doc.add_paragraph()
tr = title.add_run("Roteiro de Apresentação")
tr.bold = True; tr.font.size = Pt(24); tr.font.color.rgb = PRIMARY

sub = doc.add_paragraph()
sr = sub.add_run("TP2 Checkpoint  ·  Efeito Dominó no Mercado de Transferências")
sr.font.size = Pt(13); sr.font.color.rgb = MUTED

# bloco de info
info = doc.add_paragraph()
info.paragraph_format.space_before = Pt(8)
info.paragraph_format.space_after = Pt(12)
set_para_shading(info, "F1F5F9")
info.paragraph_format.left_indent = Cm(0.3)
for lab, val in [("Duração total: ", "5 minutos   "),
                 ("Slides: ", "10   "),
                 ("Ritmo médio: ", "~30s por slide")]:
    rl = info.add_run(lab); rl.bold = True; rl.font.color.rgb = PRIMARY
    rl.font.size = Pt(11)
    rv = info.add_run(val); rv.font.size = Pt(11); rv.font.color.rgb = DARK

# legenda
legend = doc.add_paragraph()
legend.paragraph_format.space_after = Pt(6)
r1 = legend.add_run("Convenções:  ")
r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = PRIMARY
for txt, c in [("negrito", DARK), (" = ênfase verbal  ·  ", MUTED),
               ("itálico", DARK), (" = leve pausa  ·  ", MUTED),
               ("[clica]", MUTED), (" = avança o slide  ·  ", MUTED),
               ("[gesto]", MUTED), (" = sugestão visual ao público", MUTED)]:
    r = legend.add_run(txt); r.font.size = Pt(10); r.font.color.rgb = c
    if "negrito" in txt: r.bold = True
    if "itálico" in txt: r.italic = True

doc.add_paragraph()  # respiro

# ============================== SLIDES ==============================

# ----- Slide 1 -----
slide_header(1, "Capa", "15s")
script_para([
    ("Boa tarde. Somos o ", {}),
    ("Grupo 2", {"bold": True}),
    (" e nosso projeto investiga o ", {}),
    ("efeito dominó no mercado de transferências", {"bold": True}),
    (" — especificamente, se existe um ", {}),
    ('"prêmio do vendedor"', {"bold": True}),
    (" quando um clube acaba de fazer uma grande venda e parte para uma "
     "compra na mesma janela. Eu vou apresentar nosso checkpoint de modelagem.", {}),
])
cue("[clica]")

# ----- Slide 2 -----
slide_header(2, "Pergunta de Pesquisa", "25s")
script_para([
    ("Para recapitular: a pergunta principal é — ", {}),
    ("dada uma venda feita pelo time A em uma janela, qual a probabilidade "
     "de outros clubes venderem para o time A ", {"italic": True}),
    ("acima do valor de mercado de referência", {"italic": True, "bold": True}),
    ("?", {"italic": True}),
])
script_para([
    ("A hipótese é que clubes recém-capitalizados ", {}),
    ("perdem poder de barganha", {"bold": True}),
    (" e pagam um sobrepreço. Mas tem um cuidado central no projeto: "
     "garantir que esse efeito seja ", {}),
    ("causal", {"bold": True}),
    (", e não só correlação. Clube rico vende caro e compra caro ao mesmo "
     "tempo — isso, por si só, não é prêmio do vendedor.", {}),
])
cue("[clica]")

# ----- Slide 3 -----
slide_header(3, "Estratégia de Modelagem", "50s")
script_para([
    ("Para responder isso com rigor, dividimos a modelagem em ", {}),
    ("duas etapas sequenciais", {"bold": True}),
    (".", {}),
])
cue("[gesto: aponta para o card da esquerda]")
script_para([
    ("A ", {}), ("Etapa 1", {"bold": True}),
    (", foco deste checkpoint, é um ", {}),
    ("modelo hedônico", {"bold": True}),
    (" que prevê o log do fee — o \"preço justo\" do jogador — a partir "
     "de ", {}),
    ("três blocos de variáveis", {"bold": True}),
    (". Primeiro, características do jogador: ", {}),
    ("idade, idade ao quadrado", {"bold": True}),
    (" (para capturar a curvatura biológica da carreira) e ", {}),
    ("log do market value", {"bold": True}),
    (" do Transfermarkt como proxy de habilidade. Segundo, dummies de ", {}),
    ("posição", {"bold": True}),
    (" — atacante, meia, defensor. Terceiro, controles estruturais de "
     "mercado: ", {}),
    ("dummies das sete ligas", {"bold": True}),
    (" e ", {}),
    ("efeitos fixos de temporada", {"bold": True}),
    (". ", {}),
    ("Importante:", {"italic": True}),
    (" nenhuma feature de comportamento do clube comprador entra aqui — "
     "só o que descreve o jogador e o contexto. O ", {}),
    ("resíduo", {"bold": True}),
    (" entre o preço pago e o preço previsto vira nossa variável "
     "dependente da Etapa 2.", {}),
])
cue("[gesto: aponta para o card da direita]")
script_para([
    ("A ", {}), ("Etapa 2", {"bold": True}),
    (" é a estimação causal via ", {}),
    ("Double Machine Learning", {"bold": True}),
    (". O ", {}),
    ("target Y", {"bold": True}),
    (" é o resíduo hedônico que acabei de mencionar. O ", {}),
    ("tratamento D", {"bold": True}),
    (" é o ", {}),
    ("log_revenue", {"bold": True}),
    (" — quanto o comprador faturou em vendas naquela temporada. E o "
     "vetor ", {}),
    ("W de confundidores tem 22 variáveis distribuídas em cinco blocos", {"bold": True}),
    (": ", {}),
    ("financeiro", {"bold": True}),
    (" do clube — gasto total, número de compras, balanço líquido; ", {}),
    ("estrutural do elenco", {"bold": True}),
    (" — tamanho, idade média, jogadores de seleção; ", {}),
    ("posição na rede", {"bold": True}),
    (" de transferências — PageRank, força de entrada e saída, fluxo "
     "líquido; ", {}),
    ("liga", {"bold": True}),
    (" — log do MV agregado e seis dummies; e ", {}),
    ("temporada", {"bold": True}),
    (". Tudo entra como controle no modelo do EconML.", {}),
])
script_para([
    ("Essa arquitetura segue Chernozhukov 2018 e é o padrão do estado da "
     "arte para causalidade com ML não-linear.", {}),
])
cue("[clica]")

# ----- Slide 4 -----
slide_header(4, "Resultados Comparativos", "45s")
script_para([
    ("Para a Etapa 1, comparamos ", {}),
    ("quatro modelos", {"bold": True}),
    (" com validação cruzada de 5 folds e holdout temporal — treino em "
     "2023 e 2024, teste em 2025.", {}),
])
script_para([
    ("O vencedor, em verde, foi o ", {}),
    ("Random Forest", {"bold": True}),
    (", com ", {}),
    ("R² de 75% no teste", {"bold": True}),
    (" e RMSE de 0,68 em escala logarítmica. Os outros três modelos "
     "ficaram próximos, todos acima de 70% de R², o que dá segurança de "
     "que o resultado ", {}),
    ("não depende da escolha do algoritmo", {"bold": True}),
    (".", {}),
])
script_para([
    ("Por que regressão e não classificação? Porque o target é ", {}),
    ("contínuo", {"bold": True}),
    (" — log do fee — e a literatura hedônica de valuation de jogadores "
     "reporta R² entre 0,70 e 0,80, então estamos exatamente na faixa esperada.", {}),
])
script_para([
    ("Escolhemos ", {}),
    ("RMSE, MAE e R²", {"bold": True}),
    (": RMSE penaliza erros grandes, que importam porque os fees variam "
     "de 250 mil a mais de 100 milhões de euros. MAE dá interpretação "
     "direta. E R² é a métrica padrão na literatura.", {}),
])
cue("[clica]")

# ----- Slide 5 -----
slide_header(5, "Configuração Experimental", "30s")
script_para([
    ("Para reprodutibilidade, todos os experimentos rodam com ", {}),
    ("seed 42", {"bold": True}),
    (", 14 features, e ", {}),
    ("2.079 transferências", {"bold": True}),
    (" em 3 temporadas e 7 ligas.", {}),
])
script_para([
    ("Os hiperparâmetros estão todos aqui — usamos 500 árvores nos "
     "modelos de boosting, profundidade 5, learning rate 0,05. ", {}),
    ("Importante:", {"italic": True}),
    (" esses valores ainda ", {}),
    ("não passaram por tuning", {"bold": True}),
    (" com Optuna; isso é um dos próximos passos.", {}),
])
script_para([
    ("O filtro de fee acima de 250 mil euros foi justificado empiricamente: "
     "abaixo desse valor, o desvio padrão do premium ratio cai pra menos "
     "de 0,30 — não há dinâmica real de mercado, é ruído administrativo.", {}),
])
cue("[clica]")

# ----- Slide 6 -----
slide_header(6, "Interpretabilidade (SHAP)", "30s")
script_para([
    ("Aplicamos ", {}),
    ("SHAP values", {"bold": True}),
    (" sobre o Random Forest. O resultado mostra que ", {}),
    ("log_mv", {"bold": True}),
    (" — o market value do Transfermarkt — domina o sinal preditivo, com "
     "cerca de ", {}),
    ("40% do impacto", {"bold": True}),
    (".", {}),
])
script_para([
    ("Em seguida vem o bloco ", {}),
    ("idade e idade ao quadrado", {"bold": True}),
    (", capturando a curvatura do valor por idade — o pico de valorização "
     "fica nos 20 anos. Depois, dummies de liga e posição.", {}),
])
script_para([
    ("Isso é consistente com a literatura — Peeters, 2018 — e replica a "
     "lógica do estudo de caso de de Ligt em 2019, onde SHAP decompôs "
     "77 milhões de euros entre rating do time e idade jovem. ", {}),
    ("Nosso pipeline está pronto para esse tipo de análise individual "
     "por transferência.", {"bold": True}),
])
cue("[clica]")

# ----- Slide 7 -----
slide_header(7, "Tratamento dos Confundidores", "35s")
script_para([
    ("Mapeamos ", {}),
    ("seis confundidores", {"bold": True}),
    (" que poderiam contaminar a causalidade. Esta é provavelmente a "
     "parte mais importante do trabalho.", {}),
])
cue("[gesto: percorre a tabela]")
script_para([
    ("C1, C4, C5 e C6", {"bold": True}),
    (" já estão controlados — parte na Etapa 1, via efeitos fixos de "
     "liga e temporada; parte na Etapa 2, via o vetor de confundidores "
     "W no Double ML, que inclui inclusive as features de rede como PageRank.", {}),
])
script_para([
    ("C2 e C3", {"bold": True}),
    (" — causalidade reversa e planejamento estratégico — ficam como ", {}),
    ("limitação reconhecida", {"bold": True}),
    (". Eles exigem data exata da transferência e dados de contexto "
     "esportivo (UCL, troca de técnico) que ainda não temos. Estão no "
     "roadmap de enriquecimento.", {}),
])
script_para([
    ("Sem esse mapeamento, qualquer correlação observada seria espúria. ", {}),
    ("É a disciplina causal que diferencia o trabalho de um valuation "
     "comum.", {"bold": True}),
])
cue("[clica]")

# ----- Slide 8 -----
slide_header(8, "Insights Preliminares", "35s")
script_para([
    ("Três números resumem o que os dados já estão dizendo:", {}),
])
cue("[gesto: aponta cada KPI]")
script_para([
    ("53,3% das transferências", {"bold": True, "accent": True}),
    (" apresentam sobrepreço em relação ao valor hedônico estimado — ou "
     "seja, ", {}),
    ("mais da metade", {"bold": True}),
    (" do mercado paga prêmio.", {}),
])
script_para([
    ("75,3% da variância", {"bold": True, "accent": True}),
    (" do log do fee é explicada pelo nosso modelo.", {}),
])
script_para([
    ("E a ", {}),
    ("mediana do prêmio é de 3,7%", {"bold": True, "accent": True}),
    (" — leve sobrepreço médio, mas com distribuição ", {}),
    ("aproximadamente simétrica", {"bold": True}),
    (" em torno de zero, o que valida usar o resíduo como variável "
     "dependente da Etapa 2.", {}),
])
script_para([
    ("Um ponto curioso: a temporada de 2025 mostra ", {}),
    ("correção de mercado", {"bold": True}),
    (" — prêmio médio negativo, após um 2024 inflacionado. Mas atenção: "
     "a hipótese do prêmio do vendedor ", {}),
    ("ainda está em aberto", {"bold": True}),
    (". Ela só será respondida com causalidade, na Etapa 2.", {}),
])
cue("[clica]")

# ----- Slide 9 -----
slide_header(9, "Próximos Passos", "30s")
script_para([
    ("Para a entrega final, temos ", {}),
    ("seis marcos em três semanas", {"bold": True}),
    (".", {}),
])
script_para([
    ("Semana 1", {"bold": True}),
    (": tuning de hiperparâmetros com Optuna, e enriquecimento do modelo "
     "hedônico com dados de performance — gols, xG, assistências, via FBref.", {}),
])
script_para([
    ("Semana 2", {"bold": True}),
    (": implementação completa do Double ML com EconML e cálculo das ", {}),
    ("métricas inovadoras", {"bold": True}),
    (": o Índice de Vulnerabilidade de Barganha, a sensibilidade do "
     "prêmio à centralidade na rede, e a evolução temporada a temporada "
     "como proxy do decaimento.", {}),
])
script_para([
    ("Semana 3", {"bold": True}),
    (": validação por Propensity Score Matching e robustez. ", {}),
    ("Entrega final no final do mês.", {"bold": True}),
])
cue("[clica]")

# ----- Slide 10 -----
slide_header(10, "Checklist (backup, opcional)", "10s")
script_para([
    ("Esse slide é um checklist interno de validação — todos os itens "
     "metodológicos cobertos estão referenciados aos slides anteriores. "
     "Fica como apoio durante a discussão.", {}),
])
script_para([
    ("Obrigado. Estou à disposição para perguntas.", {"italic": True}),
])

# ============================== APÊNDICE A ==============================
doc.add_page_break()
h(1, "Apêndice A — Q&A previsão (perguntas prováveis)")

qa(
    "Por que escolheram Random Forest e não XGBoost, que costuma ser mais preciso?",
    [
        ("Os quatro modelos ficaram próximos (gap < 5%). RF venceu marginalmente em "
         "RMSE de teste e tem a vantagem de OOB error, que dá uma terceira estimativa "
         "de generalização. Ainda assim, ", {}),
        ("após o tuning com Optuna, podemos reavaliar", {"bold": True}),
        (" — o ranking pode mudar.", {}),
    ],
)
qa(
    "Por que não usaram Lasso ou regressão linear como baseline?",
    [
        ("A literatura mostra que modelos lineares ", {}),
        ("subestimam o R²", {"bold": True}),
        (" em valuation de jogadores devido à forte não-linearidade da curva "
         "idade-valor. Os 4 modelos não-lineares já cobrem o espectro: árvore, "
         "boosting e kernel.", {}),
    ],
)
qa(
    "53% das transferências têm sobrepreço — isso é muito ou pouco?",
    [
        ("Por construção, o resíduo tem média próxima de zero (na verdade, −0,019). "
         "A leve assimetria (mediana de +3,7%) sugere que ", {}),
        ("o mercado paga, em média, um pequeno ágio", {"bold": True}),
        (" — coerente com a hipótese de excesso de demanda em janelas curtas. ", {}),
        ("A magnitude causal real só sai na Etapa 2.", {"bold": True}),
    ],
)
qa(
    "Como vocês evitam o vazamento de dados na Etapa 1?",
    [
        ("O modelo hedônico ", {}),
        ("não inclui nenhuma feature de comportamento do clube comprador", {"bold": True}),
        (" — só atributos do jogador e controles estruturais (liga, temporada). "
         "Variáveis como revenue_sales, total_spend e features de rede são "
         "deliberadamente ", {}),
        ("excluídas da Etapa 1", {"bold": True}),
        (" e entram apenas como confundidores na Etapa 2.", {}),
    ],
)
qa(
    "Vocês têm dados de quantas temporadas?",
    [
        ("Três — 2023, 2024 e 2025. ", {}),
        ("Reconhecemos que é uma série curta", {"bold": True}),
        (" para identificação causal com efeitos fixos por clube. A próxima "
         "iteração vai estender para 2020–2025 usando o repositório "
         "dcaribou/transfermarkt-datasets.", {}),
    ],
)
qa(
    "E se o efeito causal for nulo na Etapa 2?",
    [
        ("Esse é um ", {}),
        ("resultado válido por si só", {"bold": True}),
        (" — sustentado por rigor metodológico, mostra que a \"intuição do prêmio "
         "do vendedor\" é, na verdade, ", {}),
        ("viés de clube rico", {"bold": True}),
        (". É inclusive uma das hipóteses que estamos testando — não estamos "
         "forçando um efeito positivo.", {}),
    ],
)

# ============================== APÊNDICE B ==============================
doc.add_page_break()
h(1, "Apêndice B — Cheat sheet de números")

add_table(
    header=["Métrica", "Valor"],
    rows=[
        ["Transferências modeladas", "2.079"],
        ["Temporadas", "2023, 2024, 2025"],
        ["Ligas", "7 (Premier League, LaLiga, Serie A, Ligue 1, Liga Portugal, "
         "Jupiler Pro, Bundesliga base)"],
        ["Random Forest · Test R²", "0,7529"],
        ["Random Forest · Test RMSE", "0,6849"],
        ["Random Forest · Test MAE", "0,5094"],
        ["% com sobrepreço (resíduo > 0)", "53,3%"],
        ["Mediana do prêmio", "+3,7%"],
        ["Média do resíduo", "−0,019"],
        ["Quartil 25 / 75 do resíduo", "−34% / +32%"],
        ["Features no modelo hedônico", "14"],
        ["Confundidores em W (Etapa 2)", "22"],
        ["Filtro de fee mínimo", "€250.000"],
        ["Winsorização do premium ratio", "1–99%"],
        ["Seed", "42"],
        ["Folds de CV", "5"],
        ["Split temporal", "treino 2023-24 · teste 2025"],
    ],
    col_widths=[Cm(7.0), Cm(9.7)],
)

# ============================== APÊNDICE C ==============================
doc.add_page_break()
h(1, "Apêndice C — Dicas finais de execução")

tips = [
    ("Ensaie a transição entre slides 6 e 7", " — é onde a apresentação muda de "
     "\"métricas técnicas\" para \"argumento causal\". Marque uma respiração "
     "curta antes do \"Mapeamos seis confundidores...\"."),
    ("Não leia os números das tabelas", " — aponte para a linha do Random Forest "
     "no slide 4 e deixe o público acompanhar. Verbalize só o R² e o RMSE."),
    ("Se faltar tempo:", " corte o slide 5 (configuração) e mencione na "
     "transição \"os hiperparâmetros estão no notebook\". Mantenha 6, 7, 8 a "
     "qualquer custo — são o coração da apresentação."),
    ("Se sobrar tempo:", " abra o slide 10 e use 30s para reforçar limitações "
     "reconhecidas (sem tuning, sem performance esportiva, 3 temporadas) — "
     "passa profissionalismo."),
    ("Volume e ritmo:", " comece firme no slide 1, desacelere nos slides 3 e 7 "
     "(conteúdo conceitual denso), acelere no slide 9 (operacional)."),
    ("Encerramento:", " depois do \"Obrigado\", não fique parado. Olhe para a "
     "banca e pergunte: \"Alguma pergunta?\" — dá a sensação de domínio."),
]
for i, (head, body) in enumerate(tips, 1):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    rh = p.add_run(head); rh.bold = True; rh.font.color.rgb = PRIMARY; rh.font.size = Pt(11)
    rb = p.add_run(body); rb.font.size = Pt(11); rb.font.color.rgb = DARK


# ---------- salva ----------
out = "/Users/lucaspedras/TP-CDAF/docs/script_apresentacao.docx"
doc.save(out)
print(f"OK: {out}")
